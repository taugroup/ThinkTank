from __future__ import annotations

from pathlib import Path
from typing import Dict, List, Any
import json
import time
import io

import streamlit as st

from agent_builder import build_local_agent
from think_tank import ThinkTank
from agno.memory.v2 import UserMemory

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch
import markdown2

DB_FILE = Path("projects_db.json")
TEMPLATE_FILE = Path("scientist_templates.json")


def _docx_bytes(project: str,
                desc: str,
                scientists: List[Dict[str, str]],
                meeting: Dict[str, any]) -> bytes:
    doc = Document()
    doc.add_heading(project, level=0)
    doc.add_paragraph(desc)

    doc.add_heading("Scientists", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Title", "Expertise", "Goal", "Role"
    for s in scientists:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = (
            s["title"], s["expertise"], s["goal"], s["role"]
        )

    doc.add_heading(f"Meeting: {meeting['topic']}", level=1)
    doc.add_paragraph(f"Rounds: {meeting['rounds']}")
    doc.add_heading("Transcript", level=2)
    for m in meeting["transcript"]:
        doc.add_paragraph(f"{m['name']}: {m['content']}")

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(meeting["summary"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def _rtf_bytes(project: str,
               desc: str,
               scientists: List[Dict[str, str]],
               meeting: Dict[str, any]) -> bytes:
    """Return a very basic RTF (manual string build, no external lib)."""
    def esc(txt: str) -> str:
        return txt.replace("\\", r"\\").replace("{", r"\{").replace("}", r"\}")

    lines = [
        r"{\rtf1\ansi\deff0",
        r"\fs32\b " + esc(project) + r"\b0\par",
        r"\fs24 " + esc(desc) + r"\par\par",
        r"\b Scientists\b0\par",
    ]
    for s in scientists:
        lines.append(esc(f"{s['title']} | {s['expertise']} | {s['goal']} | {s['role']}") + r"\par")
    lines.extend([
        r"\par\b Meeting: " + esc(meeting["topic"]) + r"\b0\par",
        esc(f"Rounds: {meeting['rounds']}") + r"\par\par",
        r"\b Transcript\b0\par",
    ])
    for m in meeting["transcript"]:
        lines.append(esc(f"{m['name']}: {m['content']}") + r"\par")
    lines.extend([r"\par\b Summary\b0\par", esc(meeting["summary"]), r"\par}", ""])
    return "\n".join(lines).encode("utf-8")

def _pdf_bytes(project: str,
               desc: str,
               scientists: List[Dict[str, str]],
               meeting: Dict[str, any]) -> bytes:
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=LETTER)
    styles = getSampleStyleSheet()
    elements = []

    # Project title and description
    elements.append(Paragraph(f"<b>{project}</b>", styles['Title']))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(desc, styles['Normal']))
    elements.append(Spacer(1, 12))

    # Scientists Table
    elements.append(Paragraph("Scientists", styles['Heading1']))
    data = [["Title", "Expertise", "Goal", "Role"]]
    for s in scientists:
        row = [
            Paragraph(s["title"], styles['Normal']),
            Paragraph(s["expertise"], styles['Normal']),
            Paragraph(s["goal"], styles['Normal']),
            Paragraph(s["role"], styles['Normal']),
        ]
        data.append(row)

    # Adjust column widths to prevent overflow
    col_widths = [1.5*inch, 2.5*inch, 2.5*inch, 1.5*inch]
    table = Table(data, colWidths=col_widths, hAlign='LEFT')
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 12))

    # Meeting Topic
    elements.append(Paragraph(f"Meeting: {meeting['topic']}", styles['Heading1']))
    elements.append(Paragraph(f"Rounds: {meeting['rounds']}", styles['Normal']))
    elements.append(Spacer(1, 12))

    # Transcript with markdown parsing
    elements.append(Paragraph("Transcript", styles['Heading2']))
    for m in meeting["transcript"]:
        name = f"<b>{m['name']}:</b> "
        content_html = markdown2.markdown(m["content"])
        content_clean = content_html.replace("<p>", "").replace("</p>", "").strip()
        elements.append(Paragraph(name + content_clean, styles['Normal']))
        elements.append(Spacer(1, 6))
    elements.append(Spacer(1, 12))

    # Summary
    elements.append(Paragraph("Summary", styles['Heading2']))
    elements.append(Paragraph(meeting["summary"], styles['Normal']))

    # Build PDF
    doc.build(elements)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes

def export_meeting(project_name: str,
                   project_desc: str,
                   scientists: List[Dict[str, str]],
                   meeting: Dict[str, any]) -> Dict[str, bytes]:
    """
    Returns {'docx': …, 'rtf': …, 'pdf': …}  - each value is file bytes.
    """
    return {
        "docx": _docx_bytes(project_name, project_desc, scientists, meeting),
        "rtf": _rtf_bytes(project_name, project_desc, scientists, meeting),
        "pdf": _pdf_bytes(project_name, project_desc, scientists, meeting),
    }

def download_function(project_name: str,
                            project_desc: str,
                            project_data: Dict[str, Any],
                            meeting: Dict[str, Any]) -> None:
    
    files = export_meeting(project_name, project_desc, project_data["scientists"], meeting)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button("⬇️  DOCX", files["docx"],
                        file_name=f"{meeting['topic']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col2:
        st.download_button("⬇️  RTF", files["rtf"],
                        file_name=f"{meeting['topic']}.rtf",
                        mime="application/rtf")
    with col3:
        st.download_button("⬇️  PDF", files["pdf"],
                        file_name=f"{meeting['topic']}.pdf",
                        mime="application/pdf")

#  project DB 

def _load_projects() -> Dict[str, Any]:
    if DB_FILE.exists():
        return json.loads(DB_FILE.read_text())
    return {}


def _save_projects(data: Dict[str, Any]):
    DB_FILE.write_text(json.dumps(data, indent=2))

#  scientist templates DB

def _load_templates() -> List[Dict[str, str]]:
    if TEMPLATE_FILE.exists():
        return json.loads(TEMPLATE_FILE.read_text())
    # seed with three defaults on first run
    defaults = [
        {"title": "Immunologist", "expertise": "Immunopathology, antibody‑antigen interactions", "goal": "Guide immune‑targeting strategies", "role": "Analyse epitope selection and immune response"},
        {"title": "Machine Learning Expert", "expertise": "Deep learning, protein sequence modelling", "goal": "Develop predictive models for design", "role": "Build & chain ML models to rank candidates"},
        {"title": "Computational Biologist", "expertise": "Protein folding simulation, molecular dynamics", "goal": "Validate structural stability", "role": "Simulate docking & refine structures"},
    ]
    TEMPLATE_FILE.write_text(json.dumps(defaults, indent=2))
    return defaults


def _save_templates(templates: List[Dict[str, str]]) -> None:
    """Persist scientist templates to disk."""
    TEMPLATE_FILE.write_text(json.dumps(templates, indent=2))

def build_custom_thinktank(project_desc: str, scientists: List[Dict[str, str]]) -> ThinkTank:
    lab = ThinkTank(project_desc)
    lab.scientists.clear()
    for sd in scientists:
        lab.scientists.append(
            build_local_agent(
                name=sd["title"],
                description=f"Expertise: {sd['expertise']}. Goal: {sd['goal']}",
                role=sd["role"],
                memory=lab._memory,
                storage=lab._storage,
            )
        )
    return lab

def run_thinktank_meeting(
    project_name: str,
    project_desc: str,
    scientists: List[Dict[str, str]],
    meeting_topic: str,
    rounds: int,
    projects_db: Dict[str, Any],
):
    """Execute a team meeting and write transcript + summary back to database."""
    lab = build_custom_thinktank(project_desc, scientists)

    st.subheader(f"🧑‍🔬 Team Meeting - {meeting_topic}")

    def write(name: str, content: str):
        st.markdown(f"**—— {name} ——**")
        st.markdown(content)

    transcript: List[Dict[str, str]] = []
    def log(name: str, content: str):
        transcript.append({"name": name, "content": content})
        write(name, content)

    # PI opening
    pi_open = lab.pi.run(
        f"You are convening a team meeting. Agenda: {meeting_topic}. Share initial guidance.",
        stream=False,
    ).content
    log(lab.pi.name, pi_open)

    # Discussion rounds 
    for r in range(1, rounds + 1):
        st.markdown(f"### 🔄 Round {r}/{rounds}")
        for sci in lab.scientists:
            resp = sci.run(
                f"Context so far:\n{lab._context()}\n\nYour contribution for round {r}:",
                stream=False,
            ).content
            lab._log("scientist", sci.name, resp)
            log(sci.name, resp)

        crit = lab.critic.run(
            f"Context so far:\n{lab._context()}\n\nCritique round {r}",
            stream=False,
        ).content
        lab._log("critic", lab.critic.name, crit)
        log(lab.critic.name, crit)

        synth = lab.pi.run(
            f"Context so far:\n{lab._context()}\n\nSynthesise round {r} and pose follow‑ups.",
            stream=False,
        ).content
        lab._log("pi", lab.pi.name, synth)
        log(lab.pi.name + " (synthesis)", synth)

    # Final summary
    summary = lab.pi.run(
        f"Context so far:\n{lab._context()}\n\nProvide the final detailed meeting summary and recommendations.",
        stream=False,
    ).content
    lab._log("pi", lab.pi.name, summary)

    st.subheader("📝 Final Meeting Summary")
    st.markdown(summary)

    lab._memory.add_user_memory(memory=UserMemory(memory=summary), user_id=project_name)


    # Save to DB
    proj = projects_db.setdefault(project_name, {"description": project_desc, "scientists": scientists, "meetings": []})
    proj["description"] = project_desc
    proj["scientists"] = scientists
    proj["meetings"].append({
        "timestamp": int(time.time()),
        "topic": meeting_topic,
        "rounds": rounds,
        "transcript": transcript,
        "summary": summary,
    })
    _save_projects(projects_db)
    st.success("Meeting complete and saved 🎉")
    return proj

st.set_page_config(page_title="Think Tank", layout="wide")
st.title("🧪 Think Tank - Lab Simulator")

projects_db = _load_projects()
project_names = sorted(projects_db.keys())

# ── Project selection / creation 
st.sidebar.header("Project Manager")
proj_choice = st.sidebar.selectbox("Project", ["➕ New project"] + project_names)

if proj_choice == "➕ New project":
    project_name = st.sidebar.text_input("New project name")
    project_data = {"description": "", "scientists": [], "meetings": []}
    if not project_name:
        st.stop()
else:
    project_name = proj_choice
    project_data = projects_db[project_name]

# ── Project description --
project_desc = st.sidebar.text_area("Project description", value=project_data.get("description", ""), height=120)

# ── Scientist roster --

# Scientist templates loaded from disk 
TEMPLATES: List[Dict[str, str]] = _load_templates()

st.sidebar.subheader("Scientists")

# Template management
with st.sidebar.expander("Manage templates", expanded=False):
    # Select existing template to load
    selected_tpl_title = st.selectbox("Load template to edit", ["<new>"] + [t["title"] for t in TEMPLATES], key="tpl_select")

    # Populate fields depending on selection
    if selected_tpl_title == "<new>":
        tpl_data = {"title": "", "expertise": "", "goal": "", "role": ""}
    else:
        tpl_data = next(t for t in TEMPLATES if t["title"] == selected_tpl_title)

    new_t = st.text_input("Title", value=tpl_data["title"], key="tpl_title")
    col1, col2 = st.columns(2)
    with col1:
        new_e = st.text_area("Expertise", value=tpl_data["expertise"], height=70, key="tpl_exp")
    with col2:
        new_g = st.text_area("Goal", value=tpl_data["goal"], height=70, key="tpl_goal")
    new_r = st.text_area("Role", value=tpl_data["role"], height=70, key="tpl_role")

    # Save / update template
    if st.button("Save template") and new_t.strip():
        TEMPLATES = [t for t in TEMPLATES if t["title"] != new_t]
        TEMPLATES.append({"title": new_t, "expertise": new_e, "goal": new_g, "role": new_r})
        _save_templates(TEMPLATES)
        (st.rerun if hasattr(st, 'rerun') else st.experimental_rerun)()

    # delete templates
    del_tpl = st.multiselect("Delete templates", [t["title"] for t in TEMPLATES])
    if del_tpl and st.button("Remove selected templates"):
        TEMPLATES = [t for t in TEMPLATES if t["title"] not in del_tpl]
        _save_templates(TEMPLATES)
        (st.rerun if hasattr(st, 'rerun') else st.experimental_rerun)()

#  Load or initialise project rows  project rows 
rows = project_data.get("scientists", [])

#  Add from template --
sel_template = st.sidebar.selectbox("Add scientist from template", ["<select>"] + [t["title"] for t in TEMPLATES])
if sel_template != "<select>":
    if sel_template not in [r["title"] for r in rows]:
        rows.append(next(t for t in TEMPLATES if t["title"] == sel_template))
        _save_projects(projects_db)
        (st.rerun if hasattr(st, 'rerun') else st.experimental_rerun)()

#  Delete scientists 
if rows:
    del_choice = st.sidebar.multiselect("Delete scientist(s)", [r["title"] for r in rows])
    if del_choice and st.sidebar.button("Remove selected scientist(s)"):
        rows = [r for r in rows if r["title"] not in del_choice]
        _save_projects(projects_db)
        (st.rerun if hasattr(st, 'rerun') else st.experimental_rerun)()

#  Manual create scientist 
if st.sidebar.button("Add blank scientist"):
    rows.append({"title": f"Scientist {len(rows)+1}", "expertise": "", "goal": "", "role": ""})

#  Editable table 
rows = rows[:8]  # limit to 8
scientist_table = st.sidebar.data_editor(rows, num_rows="dynamic", use_container_width=True, key="scientist_table")

num_scientists = len(scientist_table)

# ── Meeting selection / creation -- / creation --
meetings = project_data.get("meetings", [])
meeting_labels = [f"{i+1}. {m['topic']}" for i, m in enumerate(meetings)]
meeting_choice = st.sidebar.selectbox("Meeting", ["New meeting"] + meeting_labels)

if meeting_choice == "New meeting":
    meeting_topic = st.sidebar.text_input("New meeting topic / title")
    rounds = int(st.sidebar.number_input("Rounds", min_value=1, value=3, step=1))
    run_btn = st.sidebar.button("Run Team Meeting")
    if run_btn:
        clean_scientists = [row for row in scientist_table if row["title"].strip()]
        if not clean_scientists:
            st.error("Provide at least one scientist")
            st.stop()
        proj = run_thinktank_meeting(project_name, project_desc, clean_scientists, meeting_topic, rounds, projects_db)
        download_function(project_name, project_desc, project_data, proj["meetings"][-1])
else:
    # Load existing meeting 
    sel_index = meeting_labels.index(meeting_choice)
    meeting = meetings[sel_index]
    st.markdown(f"## 🗂️ Meeting Record - {meeting['topic']}")
    for msg in meeting["transcript"]:
        st.markdown(f"**{msg['name']}**: {msg['content']}")
    st.markdown("### Summary")
    st.markdown(meeting["summary"])
    download_function(project_name, project_desc, project_data, meeting)