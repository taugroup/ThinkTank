from datetime import datetime
from typing import Dict, List
import tempfile
import re, os, copy, shutil, subprocess, tempfile
import pandas as pd
import pypandoc
pypandoc.download_pandoc()

from docx import Document
from docx.shared import Inches, Pt

import io

MERMAID_BLOCK = re.compile(r"```mermaid(.*?)```", re.S)

def _render_mermaid(code: str, out_png: str):
    """Render mermaid code to PNG with mermaid-cli (`mmdc`)."""
    with tempfile.NamedTemporaryFile("w", suffix=".mmd", delete=False) as tmp:
        tmp.write(code)
        tmp_path = tmp.name
    try:
        subprocess.run(
            ["mmdc", "-i", tmp_path, "-o", out_png, "-b", "transparent"],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
    finally:
        os.remove(tmp_path)

def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def clean_name(name: str, min_len: int = 3) -> str:
    # replace invalid chars with _
    clean = re.sub(r"[^A-Za-z0-9._-]", "_", name)
    # collapse consecutive underscores / dots / dashes
    clean = re.sub(r"[_\.-]{2,}", "_", clean)
    # strip leading / trailing non-alphanumerics
    clean = re.sub(r"^[^A-Za-z0-9]+|[^A-Za-z0-9]+$", "", clean)
    # fallback if too short
    if len(clean) < min_len:
        clean = (clean + "___")[:min_len]
    return clean[:512]

def indent(text: str, pad: int = 2) -> str:
    prefix = " " * pad
    return "\n".join(prefix + ln for ln in text.splitlines())

def _docx_bytes(project_name: str,
                project_desc: str,
                scientists: List[Dict[str, str]],
                md_text: List[str],
                table_font_size: int = 10) -> bytes:
    md_list = copy.deepcopy(md_text)        # never mutate caller’s list
    table_md = pd.DataFrame(scientists).to_markdown(index=False, tablefmt="pipe")
    md_list.insert(0, table_md)             # put table above any meeting notes
    # Add watermark at the end of list
    footer = (
        f'---\n\n')
    body_md   = "\n\n".join(md_list)
    header_md = (
        f"# {project_name}\n\n"
        f"---\n\n"
        f"## {project_desc}\n\n"
        f"**Exported on:** {now()}\n"
    )
    full_markdown = f"{header_md}\n\n{body_md}\n\n{footer}"

    images_dir = tempfile.mkdtemp()
    def _replace(match, counter=[0]):
        counter[0] += 1
        code = match.group(1).strip()
        img_path = os.path.join(images_dir, f"mermaid_{counter[0]}.png")
        _render_mermaid(code, img_path)
        return f"![diagram-{counter[0]}]({img_path})"

    full_md = MERMAID_BLOCK.sub(_replace, full_markdown)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    pypandoc.convert_text(full_md, to="docx", format="md", outputfile=tmp_path)

    doc = Document(tmp_path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(table_font_size)
    para = doc.add_paragraph()
    run = para.add_run()
    with open("assets/Logo_tau.png", "rb") as f: 
        pic_bytes = f.read()
    run.add_picture(io.BytesIO(pic_bytes), width=Inches(0.4))
    para.add_run("   Developed by TAU Group").font.size = Pt(14)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
        out_path = out.name
    doc.save(out_path)

    with open(out_path, "rb") as f:
        docx_bytes = f.read()

    os.remove(tmp_path)
    os.remove(out_path)
    shutil.rmtree(images_dir, ignore_errors=True)

    return docx_bytes

def export_meeting(project_name: str,
                   project_desc: str,
                   scientists: List[Dict[str, str]],
                   meeting: Dict[str, any],
                   md_text: List[str]) -> Dict[str, bytes]:
    """
    Returns {'docx': …, 'rtf': …, 'pdf': …}  - each value is file bytes.
    """
    return {
        "docx": _docx_bytes(project_name, project_desc, scientists, md_text)
    }